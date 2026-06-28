# -*- coding: utf-8 -*-
"""
Resume Agent — drop-in module for the AI Mock Interview app.

Pipeline:
    Upload Resume -> Resume Parser -> Skill Extraction -> ATS Analysis
        -> AI Suggestions -> Download Updated Resume

Reuses the app's Groq client and CSS classes (feedback-box, score-badge,
score-badge-low). Requires: pypdf, python-docx (for PDF/DOCX upload + download).
"""

import json
from io import BytesIO
import streamlit as st

# The label that appears in the sidebar "Assistant Mode" selectbox.
RESUME_AGENT_MODE = "📄 Resume Agent"


# ============================================================
# 1. Resume Parser — file readers
# ============================================================
def _extract_pdf(file) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file)
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(file) -> str:
    import docx

    document = docx.Document(file)
    return "\n".join(p.text for p in document.paragraphs)


def extract_resume_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return _extract_pdf(uploaded_file)
    if name.endswith(".docx"):
        return _extract_docx(uploaded_file)
    return uploaded_file.read().decode("utf-8", errors="ignore")


# ============================================================
# 2. ATS Analysis + AI Suggestions — the Groq call
# ============================================================
def _build_prompt(resume: str, target: str) -> str:
    target_block = ""
    if target.strip():
        target_block = f"TARGET ROLE / JOB DESCRIPTION:\n{target.strip()}\n\n"

    return f"""You are an expert resume reviewer and ATS (applicant tracking system) specialist. Analyze the resume below.

{target_block}RESUME:
\"\"\"
{resume.strip()}
\"\"\"

Return ONLY a valid JSON object (no markdown, no commentary, no code fences) with EXACTLY this shape:
{{
  "detectedRole": "the role this resume targets, short",
  "atsScore": <integer 0-100>,
  "verdict": "one encouraging but honest sentence about the resume's current state",
  "scoreBreakdown": [
    {{"label": "Keywords & skills", "score": <int 0-25>, "note": "short reason"}},
    {{"label": "Impact & metrics", "score": <int 0-25>, "note": "short reason"}},
    {{"label": "Formatting & ATS parse", "score": <int 0-25>, "note": "short reason"}},
    {{"label": "Clarity & structure", "score": <int 0-25>, "note": "short reason"}}
  ],
  "extractedSkills": ["up to 14 skills found in the resume"],
  "suggestedSkills": [
    {{"skill": "...", "why": "why it helps for this role, under 12 words"}}
  ],
  "professionalSummary": "a strong rewritten 2-3 sentence professional summary tailored to the role",
  "improvedBullets": [
    {{"original": "a real weak bullet from the resume", "improved": "stronger version: action verb + metric"}}
  ],
  "atsIssues": [
    {{"severity": "high|medium|low", "issue": "the problem", "fix": "the specific fix"}}
  ],
  "quickWins": ["3-5 short actionable improvements"]
}}

Rules: the four scoreBreakdown scores must sum to atsScore. Be specific to THIS resume — reference real content. Limit suggestedSkills to 6, improvedBullets to 5, atsIssues to 5. Output JSON only."""


def analyze_resume(client, model: str, resume: str, target: str) -> dict:
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "system",
                "content": "You are an expert resume reviewer and ATS specialist. You output only valid JSON, no markdown.",
            },
            {"role": "user", "content": _build_prompt(resume, target)},
        ],
        stream=False,
        temperature=0.3,
    )
    text = resp.choices[0].message.content.strip()
    text = text.replace("```json", "").replace("```", "").strip()
    start, end = text.find("{"), text.rfind("}")
    if start >= 0 and end >= 0:
        text = text[start : end + 1]
    return json.loads(text)


# ============================================================
# 3. Download Updated Resume — rewrite + .docx / .md builders
# ============================================================
def generate_updated_resume(client, model, resume, target, analysis) -> str:
    """Second pass: produce a clean, ATS-friendly rewritten resume (plain text)."""
    summary = analysis.get("professionalSummary", "")
    add_skills = ", ".join(s.get("skill", "") for s in analysis.get("suggestedSkills", []))
    bullets = "\n".join(
        f"- BEFORE: {b.get('original','')}\n  AFTER: {b.get('improved','')}"
        for b in analysis.get("improvedBullets", [])
    )
    target_block = f"TARGET ROLE: {target.strip()}\n\n" if target.strip() else ""

    prompt = f"""You are a professional resume writer. Rewrite the resume below into a clean, ATS-friendly version.

{target_block}IMPROVED SUMMARY TO USE AT THE TOP:
{summary}

SKILLS TO WEAVE IN (only where the candidate plausibly has them — never fabricate):
{add_skills}

BULLET IMPROVEMENTS TO APPLY:
{bullets}

ORIGINAL RESUME:
\"\"\"
{resume.strip()}
\"\"\"

Rules:
- Keep it 100% truthful. Do NOT invent jobs, dates, degrees, or metrics that aren't supported.
- Use clear UPPERCASE section headers (e.g. SUMMARY, SKILLS, EXPERIENCE, EDUCATION, PROJECTS).
- Use "- " for bullet points. Plain text only, no markdown symbols like ** or #.
- Keep contact details from the original if present.
Output ONLY the finished resume text."""

    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are an expert resume writer. Output clean plain-text resumes only."},
            {"role": "user", "content": prompt},
        ],
        stream=False,
        temperature=0.4,
    )
    return resp.choices[0].message.content.strip()


def _build_docx_bytes(text: str) -> bytes:
    """Turn plain-text resume into a simple, clean .docx in memory."""
    import docx
    from docx.shared import Pt

    document = docx.Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line:
            document.add_paragraph("")
            continue
        stripped = line.strip()
        # Heuristic: UPPERCASE-ish short lines are section headers
        is_header = (
            stripped == stripped.upper()
            and len(stripped) <= 40
            and any(c.isalpha() for c in stripped)
        )
        if is_header:
            h = document.add_paragraph()
            run = h.add_run(stripped)
            run.bold = True
            run.font.size = Pt(13)
        elif stripped.startswith("- ") or stripped.startswith("• "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        else:
            document.add_paragraph(stripped)

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


# ============================================================
# 4. Rendering helpers
# ============================================================
def _esc(s):
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _chips(items, bg="#1a2744", color="#90CAF9", border="#1E88E5"):
    spans = "".join(
        f'<span style="display:inline-block;background:{bg};color:{color};'
        f"border:1px solid {border};border-radius:16px;padding:4px 12px;"
        f'margin:3px;font-size:0.85rem;">{_esc(s)}</span>'
        for s in items
    )
    return f'<div style="line-height:2.2">{spans}</div>'


def _pipeline_strip(active="done"):
    """A small visual of the 6-stage pipeline."""
    stages = [
        "📥 Upload",
        "🧾 Parse",
        "🧩 Skills",
        "📊 ATS",
        "🤖 Suggestions",
        "⬇️ Download",
    ]
    html = '<div style="display:flex;flex-wrap:wrap;gap:6px;margin:6px 0 14px;">'
    for s in stages:
        html += (
            f'<span style="background:#0d2137;border:1px solid #26C6DA;color:#B2EBF2;'
            f'border-radius:6px;padding:4px 10px;font-size:0.8rem;">{s}</span>'
        )
    html += "</div>"
    return html


def _render_results(r: dict):
    score = max(0, min(100, int(r.get("atsScore", 0))))
    badge_class = "score-badge" if score >= 60 else "score-badge-low"
    verdict_label = "Strong" if score >= 75 else "Needs work" if score >= 50 else "Major gaps"

    st.markdown(
        f'<span class="{badge_class}">ATS Score: {score}/100 · {verdict_label}</span>',
        unsafe_allow_html=True,
    )
    st.markdown(
        f'<div class="feedback-box"><b>Read as:</b> {_esc(r.get("detectedRole","—"))}.<br>'
        f'{_esc(r.get("verdict",""))}</div>',
        unsafe_allow_html=True,
    )

    for b in r.get("scoreBreakdown", []):
        sc = int(b.get("score", 0))
        st.markdown(f"**{_esc(b.get('label',''))}** — {sc}/25")
        st.progress(min(1.0, sc / 25.0))
        if b.get("note"):
            st.caption(_esc(b["note"]))

    st.markdown("---")

    st.markdown("### ✅ Skills you already have")
    found = r.get("extractedSkills", [])
    st.markdown(
        _chips(found) if found else "_None detected — add a dedicated skills section._",
        unsafe_allow_html=True,
    )

    st.markdown("### ➕ Skills worth adding")
    add = r.get("suggestedSkills", [])
    if add:
        for s in add:
            st.markdown(f"- **{_esc(s.get('skill',''))}** — {_esc(s.get('why',''))}")
    else:
        st.caption("Good coverage already.")

    st.markdown("### ✍️ A sharper professional summary")
    summary = r.get("professionalSummary", "")
    st.markdown(f'<div class="feedback-box">{_esc(summary)}</div>', unsafe_allow_html=True)
    st.code(summary, language=None)

    bullets = r.get("improvedBullets", [])
    if bullets:
        st.markdown("### 🖊️ Line-by-line redlines")
        for b in bullets:
            st.markdown(
                f'<div class="feedback-box">'
                f'<span style="color:#ef9a9a;text-decoration:line-through;">{_esc(b.get("original",""))}</span><br>'
                f'<span style="color:#A5D6A7;">▸ {_esc(b.get("improved",""))}</span></div>',
                unsafe_allow_html=True,
            )

    issues = r.get("atsIssues", [])
    if issues:
        st.markdown("### 🚩 ATS flags")
        dot = {"high": "🔴", "medium": "🟠", "low": "🔵"}
        for i in issues:
            st.markdown(
                f"{dot.get(i.get('severity','low'),'🔵')} **{_esc(i.get('issue',''))}**  \n"
                f"&nbsp;&nbsp;&nbsp;*Fix:* {_esc(i.get('fix',''))}"
            )

    wins = r.get("quickWins", [])
    if wins:
        st.markdown("### ⚡ Quick wins")
        for w in wins:
            st.markdown(f"- {_esc(w)}")


# ============================================================
# 5. The main panel
# ============================================================
def render_resume_agent(
    client, model="llama-3.1-8b-instant", voice_enabled=False, speak_async=None
):
    st.markdown("## 📄 Resume Agent")
    st.markdown(_pipeline_strip(), unsafe_allow_html=True)
    st.caption(
        "Upload or paste your resume. The agent parses it, extracts skills, runs an "
        "ATS analysis, suggests improvements, and builds a downloadable updated resume."
    )

    col1, col2 = st.columns([3, 2])
    with col1:
        uploaded = st.file_uploader(
            "📥 Upload resume (PDF / DOCX / TXT)",
            type=["pdf", "docx", "txt"],
            key="resume_upload",
        )
        resume_text = st.text_area(
            "…or paste resume text",
            value=st.session_state.get("resume_text", ""),
            height=240,
            key="resume_text_area",
            placeholder="Paste your full resume here — summary, experience, skills, education.",
        )
    with col2:
        target = st.text_area(
            "🎯 Target role or job description (optional)",
            height=120,
            key="resume_target",
            placeholder="e.g. Java Backend Developer, or paste a job posting.",
        )
        st.caption("Adding a target sharpens the keyword scoring.")

    # ── Resume Parser ──
    if uploaded is not None:
        try:
            extracted = extract_resume_text(uploaded)
            if extracted.strip():
                resume_text = extracted
                st.session_state["resume_text"] = extracted
                st.success(f"✅ Parsed {uploaded.name} — {len(extracted):,} characters")
        except Exception as e:
            st.error(
                f"Couldn't read that file ({e}). Paste the text instead, or run "
                "`pip install pypdf python-docx` and restart."
            )

    # ── Run the pipeline ──
    if st.button("🚀 Run resume pipeline", type="primary", use_container_width=True):
        if len(resume_text.strip()) < 60:
            st.warning("Add a bit more — paste your resume or upload a file first.")
        else:
            # New run invalidates any previously generated updated resume
            st.session_state.pop("resume_updated_text", None)
            try:
                with st.status("Running resume pipeline…", expanded=True) as status:
                    status.update(label="🧾 Resume Parser")
                    st.write(f"✓ Parsed {len(resume_text):,} characters")

                    status.update(label="🧩 Skill Extraction · 📊 ATS Analysis · 🤖 AI Suggestions")
                    result = analyze_resume(client, model, resume_text, target)
                    st.session_state["resume_result"] = result

                    st.write(
                        f"✓ Skill Extraction — {len(result.get('extractedSkills', []))} skills found"
                    )
                    st.write(f"✓ ATS Analysis — score {result.get('atsScore', 0)}/100")
                    st.write(
                        f"✓ AI Suggestions — {len(result.get('suggestedSkills', []))} new skills, "
                        f"{len(result.get('improvedBullets', []))} bullet rewrites"
                    )
                    status.update(label="Pipeline complete ✓", state="complete", expanded=False)

                if voice_enabled and speak_async:
                    speak_async(
                        f"Your resume scored {result.get('atsScore', 0)} out of 100. "
                        + str(result.get("verdict", ""))[:240]
                    )
            except json.JSONDecodeError:
                st.error(
                    "The model returned something I couldn't parse as JSON. Try again, "
                    "or switch to a stronger model (e.g. llama-3.3-70b-versatile)."
                )
            except Exception as e:
                st.error(f"Pipeline failed: {e}")

    # ── Show results (persist across reruns) ──
    if st.session_state.get("resume_result"):
        st.markdown("---")
        _render_results(st.session_state["resume_result"])

        # ── Download Updated Resume ──
        st.markdown("---")
        st.markdown("### ⬇️ Download updated resume")
        st.caption(
            "Generates a clean, ATS-friendly rewrite that applies the summary, skills, "
            "and bullet improvements above. It stays truthful to your original."
        )

        if st.button("🪄 Generate updated resume", use_container_width=True):
            with st.spinner("Rewriting your resume…"):
                try:
                    st.session_state["resume_updated_text"] = generate_updated_resume(
                        client,
                        model,
                        st.session_state.get("resume_text", resume_text),
                        st.session_state.get("resume_target", ""),
                        st.session_state["resume_result"],
                    )
                except Exception as e:
                    st.error(f"Couldn't generate the updated resume: {e}")

        updated = st.session_state.get("resume_updated_text")
        if updated:
            st.text_area("Preview", value=updated, height=320, key="resume_updated_preview")

            dcol1, dcol2 = st.columns(2)
            with dcol1:
                try:
                    docx_bytes = _build_docx_bytes(updated)
                    st.download_button(
                        "⬇️ Download .docx",
                        data=docx_bytes,
                        file_name="updated_resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.caption(f"DOCX unavailable ({e}). Use the text download instead.")
            with dcol2:
                st.download_button(
                    "⬇️ Download .txt",
                    data=updated.encode("utf-8"),
                    file_name="updated_resume.txt",
                    mime="text/plain",
                    use_container_width=True,
                )
